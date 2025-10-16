from django.http import JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.contrib.auth.models import User
from django.contrib import auth
from .models import ChatSession, Chat, Document
from dotenv import load_dotenv
from django.utils import timezone
from openai import OpenAI
from django.core.files.storage import default_storage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document as LangchainDocument
import os
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)
embeddings = OpenAIEmbeddings()


def register(request):
    if request.method == 'POST':
        username = request.POST['username']
        email = request.POST['email']
        password1 = request.POST['password1']
        password2 = request.POST['password2']
        if password1 == password2:
            try:
                user = User.objects.create_user(username, email, password1)
                user.save()
                auth.login(request, user)
                return redirect('chatbot')
            except:
                error_message = 'Error creating user. Username may already exist.'
                return render(request, 'register.html', {'error_message': error_message})
        else:
            error_message = 'Passwords do not match'
            return render(request, 'register.html', {'error_message': error_message})
    return render(request, 'register.html')


def login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = auth.authenticate(request, username=username, password=password)
        if user is not None:
            auth.login(request, user)
            return redirect('chatbot')
        else:
            error_message = 'Invalid Username & Password'
            return render(request, 'login.html', {'error_message': error_message})
    else:
        return render(request, 'login.html')


def logout(request):
    auth.logout(request)
    return redirect('login')


def new_chat(request):
    if request.user.is_authenticated:
        new_session = ChatSession.objects.create(user=request.user)
        return redirect('chat_session', slug=new_session.slug)
    return redirect('login')


def process_file(file, filename, session):
    """
    Simple file processing function that extracts text and saves everything in database
    """
    try:
        # Get file extension
        file_extension = os.path.splitext(filename)[1].lower()
        
        # Step 1: Extract text from different file types
        text_content = ""
        
        try:
            if file_extension == ".txt":
                # Read text file
                text_content = file.read().decode("utf-8", errors="ignore")
                
            elif file_extension == ".pdf":
                # Read PDF file
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                    
            elif file_extension == ".docx":
                # Read Word document
                doc = DocxDocument(file)
                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"
            else:
                return f"Sorry, I can't read {file_extension} files. Please upload .txt, .pdf, or .docx files."
        except Exception as file_error:
            return f"❌ Error reading file: {str(file_error)}"
        
        # Step 2: Check if we got any text
        if not text_content.strip():
            return "The file appears to be empty or I couldn't extract any text from it."
        
        # Step 3: Split text into smaller chunks for better search
        try:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,      # Each chunk has max 500 characters
                chunk_overlap=50      # 50 characters overlap between chunks
            )
            text_chunks = text_splitter.split_text(text_content)
        except Exception as chunk_error:
            return f"❌ Error processing text chunks: {str(chunk_error)}"
        
        # Step 4: Create searchable index
        try:
            # Convert text chunks to documents for AI
            ai_documents = []
            for chunk in text_chunks:
                ai_documents.append(LangchainDocument(page_content=chunk))
            
            # Create searchable index using AI embeddings
            search_index = FAISS.from_documents(ai_documents, embeddings)
        except Exception as embedding_error:
            print(f"Embedding error: {embedding_error}")
            # Continue without FAISS index - will use simple text search
            search_index = None
        
        # Step 5: Save FAISS index to memory (not disk)
        import pickle
        import io
        
        # Get FAISS index data - handle potential serialization issues
        faiss_data = None
        faiss_pickle = None
        
        if search_index is not None:
            try:
                faiss_data = search_index.serialize_to_bytes()
                faiss_pickle = pickle.dumps(search_index)
            except Exception as serialization_error:
                print(f"FAISS serialization error: {serialization_error}")
                # Fallback: create a simple text-based search instead
                faiss_data = None
                faiss_pickle = None
        else:
            print("No FAISS index created, will use simple text search")
        
        # Step 6: Save everything to database
        new_document = Document.objects.create(
            session=session,
            user=session.user,
            file=file,
            filename=filename,
            file_type=file_extension.replace('.', ''),
            chunk_count=len(text_chunks),
            processed=True,
            file_content=text_content,  # Store full text content
            faiss_index_data=faiss_data,  # Store FAISS index (may be None)
            faiss_index_pickle=faiss_pickle  # Store pickled index (may be None)
        )
        
        # Step 7: Mark session as having document
        session.has_document = True
        session.save()
        
        # Return success message
        return f"✅ File '{filename}' uploaded successfully! I can now answer questions about it. ({len(text_chunks)} sections processed)"
    
    except Exception as error:
        return f"❌ Error processing file: {str(error)}"


def simple_text_search(user_question, document_content):
    """
    Simple text-based search when FAISS is not available
    """
    try:
        # Simple keyword matching
        question_words = user_question.lower().split()
        content_lower = document_content.lower()
        
        # Find sentences that contain question words
        sentences = document_content.split('.')
        relevant_sentences = []
        
        for sentence in sentences:
            sentence_lower = sentence.lower()
            # Count how many question words appear in this sentence
            word_matches = sum(1 for word in question_words if word in sentence_lower)
            if word_matches > 0:
                relevant_sentences.append((sentence.strip(), word_matches))
        
        # Sort by relevance (most matches first)
        relevant_sentences.sort(key=lambda x: x[1], reverse=True)
        
        # Take top 3 most relevant sentences
        top_sentences = [sent[0] for sent in relevant_sentences[:3]]
        
        if not top_sentences:
            return None
        
        # Create context from relevant sentences
        context = ". ".join(top_sentences)
        
        # Ask AI with simple context
        ai_prompt = f"""You are a helpful assistant. Answer the user's question based on the provided document content.

Document Content:
{context}

User Question: {user_question}

Please provide a helpful answer based on the document content above. If the answer isn't in the document, say so."""
        
        # Ask AI for answer
        ai_response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": ai_prompt}],
            temperature=0.7,
            max_tokens=500
        )
        
        return ai_response.choices[0].message.content.strip()
        
    except Exception as error:
        print(f"Error in simple text search: {error}")
        return None


def rag_answer(user_question, session):
    """
    Simple RAG function that searches document from database and answers questions
    """
    try:
        # Step 1: Get the document from database
        document = Document.objects.filter(session=session, processed=True).first()
        if not document:
            print(f"No processed document found for session {session.slug}")
            return None
        
        # Step 2: Load the search index from database
        import pickle
        import io
        
        # Load FAISS index from database
        if document.faiss_index_pickle:
            try:
                search_index = pickle.loads(document.faiss_index_pickle)
            except Exception as pickle_error:
                print(f"Error loading FAISS index: {pickle_error}")
                # Fallback to simple text search
                return simple_text_search(user_question, document.file_content)
        else:
            print("No FAISS index found in database, using simple text search")
            # Fallback to simple text search
            return simple_text_search(user_question, document.file_content)
        
        # Step 3: Search for relevant text chunks
        relevant_chunks = search_index.similarity_search(user_question, k=3)
        
        # Step 4: Check if we found any relevant content
        if not relevant_chunks:
            print("No relevant content found in document")
            return None
        
        # Step 5: Combine relevant chunks into context
        context_text = ""
        for chunk in relevant_chunks:
            context_text += chunk.page_content + "\n\n"
        
        # Step 6: Create a clear prompt for the AI
        ai_prompt = f"""You are a helpful assistant. Answer the user's question based on the provided document content.

Document Content:
{context_text}

User Question: {user_question}

Please provide a helpful answer based on the document content above. If the answer isn't in the document, say so."""
        
        # Step 7: Ask AI for answer
        ai_response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": ai_prompt}],
            temperature=0.7,
            max_tokens=500
        )
        
        # Step 8: Return the AI's answer
        return ai_response.choices[0].message.content.strip()
    
    except Exception as error:
        print(f"Error in RAG search: {error}")
        return None


def ask_ai(message, session):
    previous_chats = Chat.objects.filter(session=session).order_by('created_at')[:10]
    
    conversation = [{"role": "system", "content": "You are a helpful assistant."}]
    
    for chat in previous_chats:
        conversation.append({"role": "user", "content": chat.message})
        conversation.append({"role": "assistant", "content": chat.response})
    
    conversation.append({"role": "user", "content": message})
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=conversation,
            max_tokens=500,
            temperature=0.7,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error: {str(e)}"


def generate_chat_title(first_message):
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                "role": "system",
                "content": (
                    "You are a helpful assistant that names chat threads. "
                    "Generate a short and concise title (2–4 words max) "
                    "that summarizes the conversation based on the first user message. "
                    "Capitalize like a proper title. Do not add quotes or punctuation.")
                },
                {"role": "user", "content": first_message}
            ],
            max_tokens=20,
            temperature=0.3,
        )
        title = response.choices[0].message.content.strip().strip('"').strip("'")
        return title[:50]
    except:
        return first_message[:30] + "..." if len(first_message) > 30 else first_message


def chatbot(request, slug=None):
    """
    Main chatbot view - handles both GET (show chat) and POST (process messages/files)
    """
    # Step 1: Check if user is logged in
    if not request.user.is_authenticated:
        return redirect('login')

    # Step 2: Get all user's chat sessions for sidebar
    user_sessions = ChatSession.objects.filter(user=request.user).order_by('-updated_at')

    # Step 3: Handle current session
    current_session = None
    chats = []
    
    if slug:
        # User is viewing a specific chat session
        current_session = get_object_or_404(ChatSession, slug=slug, user=request.user)
        chats = Chat.objects.filter(session=current_session, user=request.user).order_by('created_at')

    # Step 4: Handle POST requests (new messages or file uploads)
    if request.method == 'POST':
        return handle_chat_request(request, current_session)
    
    # Step 5: Show the chat interface
    return render(request, 'chatbot.html', {
        'chats': chats,
        'user_sessions': user_sessions,
        'current_session': current_session
    })


def handle_chat_request(request, current_session):
    """
    Handle incoming chat messages and file uploads
    """
    try:
        # Check if user uploaded a file
        uploaded_file = request.FILES.get('file')
        if uploaded_file:
            return handle_file_upload(request, uploaded_file, current_session)
        
        # Handle text message
        user_message = request.POST.get('message', '').strip()
        if not user_message:
            return JsonResponse({'error': 'Please type a message'}, status=400)
        
        # Create new session if needed
        if not current_session:
            current_session = ChatSession.objects.create(user=request.user, title='New Chat')
        
        # Get AI response
        ai_response, is_from_document = get_ai_response(user_message, current_session)
        
        # Save the conversation
        Chat.objects.create(
            session=current_session,
            user=request.user,
            message=user_message,
            response=ai_response,
            is_rag_response=is_from_document,
            created_at=timezone.now()
        )
        
        # Update session title if it's still "New Chat"
        if current_session.title == 'New Chat':
            current_session.title = generate_chat_title(user_message)
        
        # Update session timestamp
        current_session.updated_at = timezone.now()
        current_session.save()
        
        return JsonResponse({
            'message': user_message,
            'response': ai_response,
            'session_slug': current_session.slug
        })
        
    except Exception as error:
        return JsonResponse({'error': f'Something went wrong: {str(error)}'}, status=500)


def handle_file_upload(request, uploaded_file, current_session):
    """
    Handle file upload and processing
    """
    try:
        # Create new session if needed
        if not current_session:
            current_session = ChatSession.objects.create(
                user=request.user,
                title=f"Document: {uploaded_file.name[:30]}"
            )
        
        # Process the file
        result_message = process_file(uploaded_file, uploaded_file.name, current_session)
        
        return JsonResponse({
            'message': 'File uploaded',
            'response': result_message,
            'session_slug': current_session.slug
        })
        
    except Exception as error:
        return JsonResponse({'error': f'File upload failed: {str(error)}'}, status=500)


def get_ai_response(user_message, session):
    """
    Get AI response - either from document (RAG) or regular chat
    """
    # Try to get answer from document first
    if session.has_document:
        document_answer = rag_answer(user_message, session)
        if document_answer:
            return document_answer, True  # Answer came from document
    
    # If no document answer, use regular AI chat
    regular_answer = ask_ai(user_message, session)
    return regular_answer, False  # Answer came from regular AI